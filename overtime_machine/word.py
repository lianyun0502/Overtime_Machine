import docx
from docx.shared import RGBColor
from docx.shared import Cm, Pt  #加入可調整的 word 單位
from docx.oxml.ns import qn
from datetime import datetime, timedelta, time
from typing import Tuple



# for idx in range(len(doc.paragraphs)):
#     if idx == 3:
#         print(doc.paragraphs[idx].text)

def write_paragraph(paragraph, row:int, text:str, font_size:int=10, font_name:str='標楷體'):
    paragraph[row].clear()
    run = paragraph[row].add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 中文字型要多加這一列
    run.font.size = Pt(font_size)

def edit_table_paragraph(cell, text:str, font_size:int=10, font_name:str='標楷體', run_idx:int=0, clear:bool=True):
    if clear:
        for paragraph in cell.paragraphs:
            paragraph.clear()
    run = cell.paragraphs[run_idx].add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 中文字型要多加這一列
    run.font.size = Pt(font_size)

def write_reason(row, col:int=0, reason:str=''):
    edit_table_paragraph(row.cells[col], reason, font_size=8)
    return

def write_predict_hour(row, col:int=1, hour:float=0.5):
    edit_table_paragraph(row.cells[col], str(hour))
    return

def write_actually_hour(row, col:int=5, hour:float=0.5):
    edit_table_paragraph(row.cells[col], str(hour))
    return

def write_overtime_type(row, col:int=6, symbol:str='v'):
    '''
    * 平日加班: col=6
    * 禮拜六加班: col=7
    * 例假日加班: col=8
    '''
    edit_table_paragraph(row.cells[col], str(symbol))
    return

def write_overtime_region(row, col:int=3, 
                          date:datetime=datetime(year=2023, month=1, day=1), 
                          region:Tuple[time]=(time(hour=10,minute=0), time(hour=19, minute=0))
                          ):
    
    start_date = f'{date.month:2d}月{date.day:2d}日'
    start_time = f'{region[0].hour:2d}時{region[0].minute:2d}分'
    end_date = f'{date.month:2d}月{date.day:2d}日'
    end_time = f'{region[1].hour:2d}時{region[1].minute:2d}分'
    for paragraph in row.cells[col].paragraphs:
        paragraph.clear()
    edit_table_paragraph(row.cells[col], start_date, run_idx=0, clear=False)
    edit_table_paragraph(row.cells[col], start_time, run_idx=1, clear=False)

    for paragraph in row.cells[col+1].paragraphs:
        paragraph.clear()
    edit_table_paragraph(row.cells[col+1], end_date, run_idx=0, clear=False)
    edit_table_paragraph(row.cells[col+1], end_time, run_idx=1, clear=False)
    return


if __name__ == '__main__':
    doc = docx.Document(r'C:\Users\eric.li\Documents\加班補休認可表-0081 Eic Li.docx')

    print('段落數量： ', len(doc.paragraphs))
    print('表格數量： ', len(doc.tables))

    Y = '111'
    date_region = f'05/21-06/20'
    e_number = '0081'
    e_name = '李璉昀 Eric Li'

    text1 = f'\t（\t{Y} ）年（{date_region}）\t  工號：\t{e_number}   姓名： {e_name}'

    write_paragraph(doc.paragraphs, 3, text1, 14)

    for table in doc.tables:
        rows = list(table.rows[2:])
        print(f'table rows number:{len(rows)}')

    for row in rows:
        write_reason(row, 0, 'Hello World,Hello World,Hello World,Hello World,Hello World')
        write_predict_hour(row, 1, 1)
        write_overtime_region(row, 3)
        write_actually_hour(row, 5, 1)
        write_overtime_type(row, 6, 'v')

    save_file_name = r'加班補休認可表-0081-2.docx'
    doc.save(save_file_name)
    print(save_file_name)

    pass